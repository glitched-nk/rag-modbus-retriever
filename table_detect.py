import re
import pandas as pd

def extract_tables_from_pdf(pdf_path: str,  page_texts: dict = None) -> list:
    import camelot
    if page_texts is None:
        page_texts = {}
    
    def _build_context(page_num: int) -> str:
        prev = page_texts.get(page_num - 1, "")
        curr = page_texts.get(page_num, "")
        combined = (prev + "\n" + curr).strip()
        return combined[:1000]
    
    tables = []
    
    #lattice
    try:
        result = camelot.read_pdf(pdf_path, pages="all", flavor="lattice")
        for t in result:
            df = t.df
            if df.shape[0] > 1 and df.shape[1] >= 2:
                page_num = t.page - 1
                tables.append({"table": df,
                               "page": page_num,
                                "context": _build_context(page_num)})
        print(f"[table_detect] lattice: {len(tables)} table(s) found")
    except Exception as e:
        print(f"[table_detect] lattice failed: {e}")

    #stream
    if not tables:
        try:
            result = camelot.read_pdf(pdf_path, pages="all", flavor="stream")
            for t in result:
                df = t.df
                if df.shape[0] > 1 and df.shape[1] >= 2:
                    page_num = t.page - 1
                    tables.append({"table": df,
                                   "page": page_num,
                                   "context": _build_context(page_num)})
            print(f"[table_detect] stream: {len(tables)} table(s) found")
        except Exception as e:
            print(f"[table_detect] stream failed: {e}")

    return tables
 
def extract_tables_from_pdfplumber(pdf_path: str, page_texts: dict = None) -> list:
    import pdfplumber
    if page_texts is None:
        page_texts = {}
    
    tables = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                prev_text = page_texts.get(page_num - 1, "")
                curr_text= page_texts.get(page_num, "")
                context = (prev_text + "\n" + curr_text).strip()[:1000]

                for t in page.extract_tables():
                    if t and len(t) > 1:
                        df = pd.DataFrame(t)
                        tables.append({"page": page.page_number,
                            "table": df,
                            "context": context })
        print(f"[table_detect] pdfplumber: {len(tables)} table(s) found")
    except Exception as e:
        print(f"[table_detect] pdfplumber failed: {e}")
    return tables

def extract_tables_from_docx(doc_path: str) -> list:
    from docx import Document
    doc = Document(doc_path)
    tables = []
    from docx.oxml.ns import qn
 
    body_children = list(doc.element.body)
    table_idx     = 0
    preceding_text_parts = []
 
    for child in body_children:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
 
        if tag == "p":
            text = "".join(node.text for node in child.iter()
                if node.tag.endswith("}t") and node.text).strip()
            if text:
                preceding_text_parts.append(text)     #Keep only the last ~10 paragraphs as context window
                if len(preceding_text_parts) > 10:
                    preceding_text_parts.pop(0)
 
        elif tag == "tbl":
            if table_idx < len(doc.tables):
                tbl = doc.tables[table_idx]
                rows = [[cell.text.strip() for cell in row.cells]
                    for row in tbl.rows]
                if rows:
                    context = "\n".join(preceding_text_parts[-10:])[:1000]
                    tables.append({
                        "table":   pd.DataFrame(rows),
                        "page":    None,
                        "context": context,
                    })
                table_idx += 1
            preceding_text_parts = []
    return tables

KEYWORDS = [
    "address", "register", "parameter", "parameters",
    "datatype", "data type", "data types",
    "float", "int32", "int16", "uint16",
    "description", "unit", "function code",
]
 
def _score_row(row_text: str) -> int:       #Count how many Modbus keywords appear in a string
    lower = row_text.lower()
    return sum(1 for kw in KEYWORDS if kw in lower)

def is_modbus_table(df: pd.DataFrame) -> bool:          #keyword check for first 5 rows
    score = 0
    for i in range(min(5, len(df))):
        row_text = " ".join(str(v) for v in df.iloc[i].tolist())
        score += _score_row(row_text)
        if score >= 2:
            break
 
    if score < 2:
        return False
 
    # Address column sanity check: look for at least 3 rows with numeric-looking addresses
    address_pattern = re.compile(r"^\d{3,7}$")
    for col_idx in range(df.shape[1]):
        col = df.iloc[:, col_idx].astype(str)
        matches = col.str.strip().apply(lambda v: bool(address_pattern.match(v))).sum()
        if matches >= 3:
            return True
    return False

COLUMN_HINTS = {
    "address": ["address", "addr", "reg", "register no", "register", "register address", "modbus address", "modbus reg"],
    "description": ["parameter", "description", "parameters", "name", "register name", "object name", "label", "variable",
                    "details", "meaning", "function"],
    "datatype": ["data type", "datatype", "type", "format", "data format"],
    "serial": ["sl.no", "sl no", "no.", "no", "s.no", "sr.no", "index"],
    "unit": ["unit", "units"],
    "scaling": ["scaling", "scale", "factor", "multiplier"],
}
 
def detect_modbus_columns(df: pd.DataFrame) -> dict:
    header_rows = min(5, len(df))
    col_texts = {}
    for col_idx in range(df.shape[1]):
        combined = " ".join(str(df.iloc[r, col_idx]) for r in range(header_rows)    ).lower()
        col_texts[col_idx] = combined
 
    roles = {role: -1 for role in COLUMN_HINTS}
 
    for role, hints in COLUMN_HINTS.items():
        for col_idx, text in col_texts.items():
            if any(hint in text for hint in hints):
                roles[role] = col_idx
                break

    # Fallback: if address still not found, look for the first column whose DATA rows contain 4-5 digit numbers
    if roles["address"] == -1:
        addr_pat = re.compile(r"^\d{3,6}$")
        for col_idx in range(df.shape[1]):
            col = df.iloc[min(5, len(df)):, col_idx].astype(str).str.strip()
            if col.apply(lambda v: bool(addr_pat.match(v))).sum() >= 3:
                roles["address"] = col_idx
                break
    if roles["description"] == -1:      #check if necessary
        assigned = {v for v in roles.values() if v != -1}
        best_col, best_len = -1, 0
        data_start = find_data_start_row(df, roles)
        for col_idx in range(df.shape[1]):
            if col_idx in assigned:
                continue
            avg = df.iloc[data_start:, col_idx].astype(str).str.strip().str.len().mean()
            if avg and avg > best_len:
                best_len, best_col = avg, col_idx
        if best_col != -1:
            roles["description"] = best_col
            print(f"    [detect_columns] description fallback → col {best_col} ")
    return roles
 
def find_data_start_row(df: pd.DataFrame, roles: dict) -> int:
    addr_col = roles.get("address", -1)
    if addr_col == -1:
        return 1  # best guess
    addr_pat = re.compile(r"^\d{3,6}$")
    for row_idx in range(len(df)):
        val = str(df.iloc[row_idx, addr_col]).strip()
        if addr_pat.match(val):
            return row_idx
    return 1