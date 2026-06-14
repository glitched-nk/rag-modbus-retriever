import os
import fitz
import pdfplumber
import pandas as pd
from docx import Document
from PIL import Image

#Return all text from a PDF using pdfplumber.
#Returns an empty string if the PDF has no text layer (scanned/image PDF)
def extract_pdf_text(pdf_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        print(f"[extractor] pdfplumber error on {pdf_path}: {e}")
    return text

def extract_pdf_pages_with_text(pdf_path: str) -> dict:
    page_texts = {}
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                page_texts[i] = t if t else ""
    except Exception as e:
        print(f"[extractor] page text extraction error: {e}")
    return page_texts

#Higher DPI (200) improves OCR accuracy on small table text.
def convert_pdf_to_images(pdf_path: str, output_folder: str, dpi: int = 200) -> list:

    os.makedirs(output_folder, exist_ok=True)
    doc = fitz.open(pdf_path)
    image_paths: list = []
    base = os.path.splitext(os.path.basename(pdf_path))[0]
 
    try:
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            mat  = fitz.Matrix(dpi / 72, dpi / 72)
            pix  = page.get_pixmap(matrix=mat)
            image_path = os.path.join(
                output_folder, f"{base}_page_{page_num:04d}.png")
            pix.save(image_path)
            image_paths.append(image_path)
    finally:
        doc.close()
    return image_paths

def extract_pdf_text_with_ocr_fallback(pdf_path: str, image_folder: str, ocr_func, min_text_threshold: int = 200) -> str:
    text = extract_pdf_text(pdf_path)

    # Enough embedded text exists
    if len(text.strip()) >= min_text_threshold:
        return text

    print(f"[extractor] Low text PDF detected -> OCR fallback")
    ocr_text_parts = []

    try:
        image_paths = convert_pdf_to_images(pdf_path, image_folder)
        # OCR only first few pages for identification
        for img_path in image_paths:
            page_text = ocr_func(img_path)
            if page_text.strip():
                ocr_text_parts.append(page_text)

    except Exception as e:
        print(f"[extractor] OCR fallback failed: {e}")

    final_text = "\n".join(ocr_text_parts)

    return final_text

#Heuristic: if pdfplumber extracts fewer than 50 characters from the first 3 pages the PDF is likely a scanned image
def is_scanned_pdf(pdf_path: str) -> bool:
    sample = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages[:3]:
                t = page.extract_text()
                if t:
                    sample += t
    except Exception:
        pass
    return len(sample.strip()) < 50


def extract_docx_text(doc_path: str) -> str:
    text_parts = []
    try:
        doc = Document(doc_path)

        # Paragraphs
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt:
                text_parts.append(txt)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )

                if row_text:
                    text_parts.append(row_text)

    except Exception as e:
        print(f"[extractor] DOCX text extraction failed: {e}")

    return "\n".join(text_parts)

def extract_excel_text(file_path: str) -> str:
    text_parts = []
    try:
        xl = pd.ExcelFile(file_path)

        for sheet_name in xl.sheet_names:

            df = pd.read_excel(file_path, sheet_name=sheet_name,header=None)
            text_parts.append(f"\n--- SHEET: {sheet_name} ---\n")

            for _, row in df.iterrows():
                row_text = " | ".join(
                    str(v).strip()
                    for v in row.tolist()
                    if str(v).strip() not in ("nan", "")
                )

                if row_text:
                    text_parts.append(row_text)

    except Exception as e:
        print(f"[extractor] Excel text extraction failed: {e}")

    return "\n".join(text_parts)

#Read all sheets from an Excel file.
#Returns {sheet_name: DataFrame}.
def extract_excel(file_path: str) -> dict:
    xl = pd.ExcelFile(file_path)
    sheets = {}
    for sheet in xl.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet, header=None)
        sheets[sheet] = df
    return sheets

def enrich_tables_with_gap_text(tables: list, page_texts: dict | None = None) -> list:      #to know text in b/w 2 tables
    if not tables:
        return tables
 
    enriched = []
 
    for i, tbl in enumerate(tables):
        tbl = dict(tbl)            # shallow copy — don't mutate the original
 
        if i == 0:
            tbl["gap_text_before"] = ""
            enriched.append(tbl)
            continue
 
        prev = enriched[i - 1]
        gap  = ""
 
        page_a = prev.get("page")
        page_b = tbl.get("page")
 
        if page_texts is not None and page_a is not None and page_b is not None:
            if page_a == page_b:
                # Both tables are on the same page: extract the text between the end of prev's context and the start of this context.
                full_page = page_texts.get(page_a, "")
                ctx_prev  = prev.get("context", "")
                ctx_curr  = tbl.get("context", "")
 
                # Find where the previous table's context ends in the page
                pos_prev_end = -1
                if ctx_prev and ctx_prev in full_page:
                    pos_prev_end = full_page.find(ctx_prev) + len(ctx_prev)
 
                # Find where this table's context starts
                pos_curr_start = -1
                if ctx_curr and ctx_curr in full_page:
                    pos_curr_start = full_page.find(ctx_curr)
 
                if pos_prev_end != -1 and pos_curr_start != -1 \
                        and pos_curr_start > pos_prev_end:
                    gap = full_page[pos_prev_end:pos_curr_start]
                else:
                    gap = ""
 
            else:
                # Tables on different pages
                parts = []
                for p in range(page_a + 1, page_b):
                    parts.append(page_texts.get(p, ""))
                # Also add text before this table's context on its own page
                ctx_curr  = tbl.get("context", "")
                full_page = page_texts.get(page_b, "")
                if ctx_curr and ctx_curr in full_page:
                    parts.append(full_page[:full_page.find(ctx_curr)])
                gap = "\n".join(parts)
 
        else:
            gap = ""
 
        tbl["gap_text_before"] = gap
        enriched.append(tbl)
 
    return enriched